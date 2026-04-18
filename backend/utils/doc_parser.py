import os
import re
from typing import List
from pathlib import Path


def parse_txt(file_path: str) -> List[str]:
    """解析TXT文件"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    # 按段落分割
    chunks = [p.strip() for p in content.split('\n\n') if p.strip()]
    return chunks


def parse_word(file_path: str) -> List[str]:
    """解析Word文档"""
    from docx import Document
    doc = Document(file_path)
    chunks = []
    current_chunk = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
        else:
            current_chunk.append(text)
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(row_text)
    return [c for c in chunks if c.strip()]


def parse_excel(file_path: str) -> List[str]:
    """解析Excel文件，每行转为一条知识"""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    chunks = []
    for sheet in wb.worksheets:
        headers = []
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            if i == 0:
                headers = row_data
                continue
            if any(row_data):
                if headers:
                    parts = [f"{headers[j]}: {row_data[j]}" for j in range(min(len(headers), len(row_data))) if row_data[j]]
                    chunks.append(' | '.join(parts))
                else:
                    chunks.append(' | '.join(filter(None, row_data)))
    wb.close()
    return [c for c in chunks if c.strip()]


def parse_document(file_path: str, file_type: str) -> List[str]:
    """根据文件类型解析文档"""
    ext = file_type.lower()
    if ext in ('txt',):
        return parse_txt(file_path)
    elif ext in ('docx', 'doc', 'word'):
        return parse_word(file_path)
    elif ext in ('xlsx', 'xls', 'excel'):
        return parse_excel(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """将长文本切片"""
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks
